"""Editable Word exam export with native OMML equations.

The application owns the document structure and layout. Pandoc is used only as
an offline LaTeX-math-to-OMML adapter; it never receives the complete exam
document. Formulas that Pandoc cannot translate are rendered as visible PNG
fallbacks instead of being silently dropped.
"""

from __future__ import annotations

import copy
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
import zipfile
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor

from mathbank.runtime_components import find_usable_pandoc


MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
TYPE_LABELS = {
    "single_choice": "选择题",
    "multi_choice": "多选题",
    "fill_in_blank": "填空题",
    "detailed_answer": "解答题",
}
TYPE_ORDER = ("single_choice", "multi_choice", "fill_in_blank", "detailed_answer")
EXAM_19_STARTS = {
    "single_choice": 1,
    "multi_choice": 9,
    "fill_in_blank": 12,
    "detailed_answer": 15,
}
IMAGE_PATTERN = re.compile(r"!\[.*?\]\(([^)]+)\)")
CHOICES_PATTERN = re.compile(
    r"\\begin\{choices\}([\s\S]*?)\\end\{choices\}", re.IGNORECASE
)
TABULAR_PATTERN = re.compile(
    r"\\begin\{tabular\}\s*\{([^}]*)\}([\s\S]*?)\\end\{tabular\}",
    re.IGNORECASE,
)
BARE_MATH_ENVS = ("cases", "aligned", "alignedat", "matrix", "pmatrix", "bmatrix", "vmatrix", "Vmatrix")
CJK_BODY_FONT = "SimSun"
CJK_HEADING_FONT = "SimHei"
CJK_TITLE_FONT = "STZhongsong"
LATIN_FONT = "Times New Roman"
STIX_TWO_MATH_PATHS = (
    Path("/System/Library/Fonts/Supplemental/STIXTwoMath.otf"),
    Path("/Library/Fonts/STIXTwoMath.otf"),
    Path.home() / "Library/Fonts/STIXTwoMath.otf",
)


def _resolve_word_math_font() -> str:
    """Choose a complete OpenType math font that exists on this platform."""
    if any(path.is_file() for path in STIX_TWO_MATH_PATHS):
        return "STIX Two Math"
    return "Cambria Math"


MATH_FONT = _resolve_word_math_font()
BODY_FONT_SIZE = 10.5
SMALL_FONT_SIZE = 10.5
SECTION_FONT_SIZE = 11.0


class WordExportError(RuntimeError):
    """Raised when a DOCX cannot be generated safely."""


@dataclass
class WordExportDiagnostics:
    native_formulas: int = 0
    fallback_formulas: int = 0
    failed_formulas: int = 0
    missing_images: int = 0
    answer_card_omitted: bool = False
    warnings: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "native_formulas": self.native_formulas,
            "fallback_formulas": self.fallback_formulas,
            "failed_formulas": self.failed_formulas,
            "missing_images": self.missing_images,
            "answer_card_omitted": self.answer_card_omitted,
            "warnings": list(self.warnings),
        }


@dataclass(frozen=True)
class MathToken:
    latex: str
    display: bool = False


@dataclass(frozen=True)
class BlankToken:
    """A Word-native underlined answer blank, kept outside OMML."""


@dataclass
class PreparedQuestion:
    question: dict
    score: int
    stem: str
    choices: list[str]
    answer: str
    images: list[Path]
    solution_space: float


def _find_pandoc() -> str | None:
    resolved = find_usable_pandoc()
    return os.fspath(resolved[0]) if resolved else None


def _normalize_formula(formula: str) -> str:
    value = (formula or "").strip()
    value = re.sub(r"^\$\$?|\$\$?$", "", value).strip()
    value = re.sub(r"\\lt\s*", "<", value)
    value = re.sub(r"\\gt\s*", ">", value)
    value = re.sub(r"\\(?:displaystyle|textstyle|scriptstyle|scriptscriptstyle)\b", "", value)
    return value.strip()


def _consume_math(text: str, start: int) -> tuple[MathToken | None, int]:
    if text.startswith("$$", start):
        end = text.find("$$", start + 2)
        if end >= 0:
            return MathToken(_normalize_formula(text[start + 2 : end]), True), end + 2
    if text.startswith(r"\[", start):
        end = text.find(r"\]", start + 2)
        if end >= 0:
            return MathToken(_normalize_formula(text[start + 2 : end]), True), end + 2
    if text.startswith(r"\(", start):
        end = text.find(r"\)", start + 2)
        if end >= 0:
            return MathToken(_normalize_formula(text[start + 2 : end]), False), end + 2
    if text[start : start + 1] == "$" and (start == 0 or text[start - 1] != "\\"):
        cursor = start + 1
        while cursor < len(text):
            if text[cursor] == "$" and text[cursor - 1] != "\\":
                return MathToken(_normalize_formula(text[start + 1 : cursor]), False), cursor + 1
            cursor += 1
    for env_name in BARE_MATH_ENVS:
        opener = rf"\begin{{{env_name}}}"
        if text.startswith(opener, start):
            closer = rf"\end{{{env_name}}}"
            end = text.find(closer, start + len(opener))
            if end >= 0:
                finish = end + len(closer)
                return MathToken(_normalize_formula(text[start:finish]), True), finish
    return None, start


def tokenize_mixed_content(text: str) -> list[str | MathToken | BlankToken]:
    """Split mixed Chinese/LaTeX text without relying on regex lookbehind."""
    source = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    tokens: list[str | MathToken | BlankToken] = []
    plain: list[str] = []
    index = 0
    while index < len(source):
        token, next_index = _consume_math(source, index)
        if token is not None:
            if plain:
                tokens.append("".join(plain))
                plain = []
            if token.latex:
                formula_parts = token.latex.split(r"\fillin")
                for part_index, formula_part in enumerate(formula_parts):
                    if formula_part.strip():
                        tokens.append(MathToken(formula_part.strip(), token.display))
                    if part_index < len(formula_parts) - 1:
                        tokens.append(BlankToken())
            index = next_index
            continue
        if source.startswith(r"\fillin", index):
            if plain:
                tokens.append("".join(plain))
                plain = []
            tokens.append(BlankToken())
            index += len(r"\fillin")
            continue
        plain.append(source[index])
        index += 1
    if plain:
        tokens.append("".join(plain))
    return tokens


def _extract_formulas(texts: Iterable[str]) -> list[str]:
    values: list[str] = []
    seen: set[str] = set()
    for text in texts:
        for token in tokenize_mixed_content(text):
            if isinstance(token, MathToken) and token.latex and token.latex not in seen:
                values.append(token.latex)
                seen.add(token.latex)
    return values


class PandocOmmlConverter:
    """Batch-convert unique LaTeX formulas and return serialized m:oMath nodes."""

    def __init__(self, pandoc_path: str | None = None):
        self.pandoc_path = pandoc_path or _find_pandoc()

    def convert_many(self, formulas: Iterable[str]) -> dict[str, bytes]:
        unique = list(dict.fromkeys(item for item in formulas if item))
        if not unique or not self.pandoc_path:
            return {}
        converted: dict[str, bytes] = {}
        self._convert_chunk(unique, converted)
        return converted

    def _convert_chunk(self, formulas: list[str], output: dict[str, bytes]) -> None:
        if not formulas:
            return
        result = self._run_pandoc(formulas)
        output.update(result)
        missing = [item for item in formulas if item not in result]
        if not missing or len(formulas) == 1:
            return
        midpoint = max(1, len(missing) // 2)
        self._convert_chunk(missing[:midpoint], output)
        self._convert_chunk(missing[midpoint:], output)

    def _run_pandoc(self, formulas: list[str]) -> dict[str, bytes]:
        markdown_parts: list[str] = []
        markers: dict[str, str] = {}
        for index, formula in enumerate(formulas):
            marker = f"MBFORMULA{index:06d}"
            markers[marker] = formula
            markdown_parts.extend((marker, "", f"$${formula}$$", ""))
        with tempfile.TemporaryDirectory(prefix="mathbank_word_math_") as temp_dir:
            temp_path = Path(temp_dir)
            source_path = temp_path / "formulas.md"
            docx_path = temp_path / "formulas.docx"
            source_path.write_text("\n".join(markdown_parts), encoding="utf-8")
            try:
                completed = subprocess.run(
                    [
                        str(self.pandoc_path),
                        "-f",
                        "markdown+tex_math_dollars",
                        "-t",
                        "docx",
                        str(source_path),
                        "-o",
                        str(docx_path),
                    ],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=45,
                    check=False,
                )
            except (OSError, subprocess.TimeoutExpired):
                return {}
            if completed.returncode != 0 or not docx_path.exists():
                return {}
            with zipfile.ZipFile(docx_path, "r") as archive:
                xml_data = archive.read("word/document.xml")

        root = parse_xml(xml_data)
        paragraphs = root.xpath(".//w:body/w:p")
        result: dict[str, bytes] = {}
        pending_formula: str | None = None
        for paragraph in paragraphs:
            plain_text = "".join(paragraph.xpath(".//w:t/text()"))
            if plain_text in markers:
                pending_formula = markers[plain_text]
                continue
            if pending_formula:
                math_nodes = paragraph.xpath(".//m:oMath")
                if math_nodes:
                    from lxml import etree

                    result[pending_formula] = etree.tostring(math_nodes[0])
                    pending_formula = None
                elif plain_text.strip():
                    pending_formula = None
        return result


def _set_cell_margins(cell, top=70, start=70, bottom=70, end=70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table, color: str | None = None, size: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil" if color is None else "single")
        if color:
            tag.set(qn("w:color"), color)
            tag.set(qn("w:sz"), str(size))


def _set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def _set_rfonts(r_pr, cjk_font: str, latin_font: str) -> None:
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), latin_font)
    r_fonts.set(qn("w:hAnsi"), latin_font)
    r_fonts.set(qn("w:cs"), latin_font)
    r_fonts.set(qn("w:eastAsia"), cjk_font)


def _set_run_font(
    run,
    size: float = BODY_FONT_SIZE,
    bold: bool = False,
    color: RGBColor | None = None,
    cjk_font: str = CJK_BODY_FONT,
    latin_font: str = LATIN_FONT,
) -> None:
    run.font.name = latin_font
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    _set_rfonts(r_pr, cjk_font, latin_font)


def _clean_plain_text(text: str) -> str:
    value = text or ""
    value = IMAGE_PATTERN.sub("", value)
    value = re.sub(r"\\begin\{tikzpicture\}[\s\S]*?\\end\{tikzpicture\}", "", value)
    value = re.sub(r"\\begin\{(?:center|table|figure)\}", "", value)
    value = re.sub(r"\\end\{(?:center|table|figure)\}", "", value)
    value = re.sub(r"\\centering\b", "", value)
    value = re.sub(r"\\(?:textbf|mathbf|mathrm|operatorname)\{([^{}]*)\}", r"\1", value)
    value = re.sub(r"\\(?:quad|qquad)\b|\\[,;!]", " ", value)
    value = value.replace(r"\%", "%").replace(r"\_", "_")
    value = value.replace("\\\\", "\n")
    return value


def _omml_character_style(char: str, previous: str = "p") -> str:
    """Return Word's semantic math style for one visible character."""
    category = unicodedata.category(char)
    if category.startswith("M"):
        return previous
    if category.startswith("L"):
        name = unicodedata.name(char, "")
        if "LATIN" in name or "GREEK" in name or "MATHEMATICAL" in name:
            return "i"
    return "p"


def _set_omml_variable_font(math_run, style: str) -> None:
    """Apply the exam's Times New Roman variable face to one OMML run."""
    if style not in {"i", "bi"}:
        return
    math_r_pr = math_run.find(qn("m:rPr"))
    if math_r_pr is not None and math_r_pr.find(qn("m:scr")) is not None:
        return
    if math_r_pr is None:
        math_r_pr = OxmlElement("m:rPr")
        math_run.insert(0, math_r_pr)
    if math_r_pr.find(qn("m:nor")) is None:
        normal_text = OxmlElement("m:nor")
        style_node = math_r_pr.find(qn("m:sty"))
        if style_node is None:
            math_r_pr.append(normal_text)
        else:
            math_r_pr.insert(math_r_pr.index(style_node), normal_text)
    word_r_pr = math_run.find(qn("w:rPr"))
    if word_r_pr is None:
        word_r_pr = OxmlElement("w:rPr")
        insertion_index = math_run.index(math_r_pr) + 1
        math_run.insert(insertion_index, word_r_pr)
    run_fonts = word_r_pr.find(qn("w:rFonts"))
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        word_r_pr.insert(0, run_fonts)
    for attribute in ("w:ascii", "w:hAnsi", "w:cs"):
        run_fonts.set(qn(attribute), LATIN_FONT)
    if word_r_pr.find(qn("w:i")) is None:
        word_r_pr.append(OxmlElement("w:i"))
    if word_r_pr.find(qn("w:iCs")) is None:
        word_r_pr.append(OxmlElement("w:iCs"))
    if style == "bi" and word_r_pr.find(qn("w:b")) is None:
        word_r_pr.append(OxmlElement("w:b"))


def _set_native_omml_run_styles(omml) -> None:
    """Make Word's plain/italic math semantics explicit without setting fonts.

    Pandoc leaves letters and digits without ``m:sty`` because Microsoft Word
    infers their usual math style.  WPS for Mac instead treats every unmarked
    run as italic, so even numerals appear slanted.  Split mixed runs when
    necessary and mark Latin/Greek variables italic while keeping numerals,
    punctuation, operators and CJK text plain.  Existing Pandoc styles (for
    operators, ``\\mathrm``, functions and special alphabets) are preserved.
    """
    for math_run in list(omml.iter(qn("m:r"))):
        math_r_pr = math_run.find(qn("m:rPr"))
        if math_r_pr is not None and len(math_r_pr):
            style_node = math_r_pr.find(qn("m:sty"))
            if style_node is not None:
                _set_omml_variable_font(math_run, style_node.get(qn("m:val"), ""))
            continue
        text_node = math_run.find(qn("m:t"))
        if text_node is None or not text_node.text:
            continue

        groups: list[tuple[str, str]] = []
        previous_style = "p"
        for char in text_node.text:
            style = _omml_character_style(char, previous_style)
            previous_style = style
            if groups and groups[-1][0] == style:
                groups[-1] = (style, groups[-1][1] + char)
            else:
                groups.append((style, char))

        parent = math_run.getparent()
        if parent is None:
            continue
        insertion_index = parent.index(math_run)
        for offset, (style, value) in enumerate(groups):
            styled_run = copy.deepcopy(math_run)
            styled_r_pr = styled_run.find(qn("m:rPr"))
            if styled_r_pr is None:
                styled_r_pr = OxmlElement("m:rPr")
                styled_run.insert(0, styled_r_pr)
            style_node = OxmlElement("m:sty")
            style_node.set(qn("m:val"), style)
            styled_r_pr.append(style_node)
            if style == "i":
                # Traditional Chinese exam papers conventionally use Times
                # New Roman Italic for Latin/Greek variables.  Restrict the
                # direct font to variable runs only: applying it to numerals
                # caused WPS for Mac to hide choice contents until equations
                # were manually rebuilt as "Professional" format.
                _set_omml_variable_font(styled_run, style)
            styled_run.find(qn("m:t")).text = value
            parent.insert(insertion_index + offset, styled_run)
        parent.remove(math_run)


def _append_omml(paragraph, xml_bytes: bytes) -> None:
    """Append Pandoc's native OMML without rewriting individual math runs.

    Word for Mac can hide otherwise valid digits and symbols when a converter
    adds direct ``w:rPr`` formatting to every ``m:r``.  Keeping Pandoc's
    minimal OMML lets Word apply the document math font and paragraph size in
    the same way as an equation created through Word's own equation editor.
    """
    omml = copy.deepcopy(parse_xml(xml_bytes))
    _set_native_omml_run_styles(omml)
    paragraph._p.append(omml)


def _formula_fallback_png(formula: str) -> bytes | None:
    xelatex = shutil.which("xelatex")
    if not xelatex:
        return None
    tex = rf"""\documentclass[preview,border=2pt]{{standalone}}
\usepackage{{amsmath,amssymb,mathtools,cancel,extarrows,yhmath}}
\begin{{document}}
$\displaystyle {formula}$
\end{{document}}
"""
    with tempfile.TemporaryDirectory(prefix="mathbank_word_fallback_") as temp_dir:
        temp_path = Path(temp_dir)
        tex_path = temp_path / "formula.tex"
        tex_path.write_text(tex, encoding="utf-8")
        try:
            completed = subprocess.run(
                [xelatex, "-interaction=nonstopmode", "-halt-on-error", tex_path.name],
                cwd=temp_path,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                timeout=30,
                check=False,
            )
            pdf_path = temp_path / "formula.pdf"
            if completed.returncode != 0 or not pdf_path.exists():
                return None
            import pymupdf as fitz
            from PIL import Image, ImageChops

            with fitz.open(pdf_path) as pdf:
                pixmap = pdf[0].get_pixmap(matrix=fitz.Matrix(3, 3), alpha=False)
                image = Image.open(BytesIO(pixmap.tobytes("png"))).convert("RGB")
            background = Image.new("RGB", image.size, "white")
            bbox = ImageChops.difference(image, background).getbbox()
            if bbox:
                left, top, right, bottom = bbox
                image = image.crop((max(0, left - 8), max(0, top - 6), min(image.width, right + 8), min(image.height, bottom + 6)))
            output = BytesIO()
            image.save(output, format="PNG")
            return output.getvalue()
        except Exception:
            return None


class WordExamBuilder:
    def __init__(self, omml_map: dict[str, bytes], diagnostics: WordExportDiagnostics):
        self.omml_map = omml_map
        self.diagnostics = diagnostics
        self.math_font = MATH_FONT
        self.fallback_cache: dict[str, bytes | None] = {}
        self.doc = Document()
        self._configure_document()

    def _configure_document(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        # Match the 1-inch geometry used by the LaTeX exam template.
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)

        normal = self.doc.styles["Normal"]
        normal.font.name = LATIN_FONT
        normal.font.size = Pt(BODY_FONT_SIZE)
        _set_rfonts(normal._element.get_or_add_rPr(), CJK_BODY_FONT, LATIN_FONT)
        # Fractions and radicals may exceed a fixed line box. "At least" keeps
        # the normal rhythm while allowing Word to expand formula-heavy lines.
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        normal.paragraph_format.line_spacing = Pt(18)
        normal.paragraph_format.space_after = Pt(0)

        for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2"):
            style = self.doc.styles[style_name]
            style.font.name = LATIN_FONT
            _set_rfonts(style._element.get_or_add_rPr(), CJK_HEADING_FONT, LATIN_FONT)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("第 ")
        _set_run_font(run, 9)
        self._append_field(footer, "PAGE")
        run = footer.add_run(" 页")
        _set_run_font(run, 9)

        settings = self.doc.settings.element
        theme_font_lang = settings.find(qn("w:themeFontLang"))
        if theme_font_lang is None:
            theme_font_lang = OxmlElement("w:themeFontLang")
            settings.append(theme_font_lang)
        theme_font_lang.set(qn("w:val"), "zh-CN")
        theme_font_lang.set(qn("w:eastAsia"), "zh-CN")

        math_pr = settings.find(qn("m:mathPr"))
        if math_pr is None:
            math_pr = OxmlElement("m:mathPr")
            settings.insert(0, math_pr)
        math_font = math_pr.find(qn("m:mathFont"))
        if math_font is None:
            math_font = OxmlElement("m:mathFont")
            math_pr.insert(0, math_font)
        math_font.set(qn("m:val"), self.math_font)

        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)

    @staticmethod
    def _append_field(paragraph, field_name: str) -> None:
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = f" {field_name} "
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        text_node = OxmlElement("w:t")
        text_node.text = "1"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for node in (begin, instruction, separate, text_node, end):
            run = OxmlElement("w:r")
            run.append(node)
            paragraph._p.append(run)

    def add_mixed(self, paragraph, text: str, size: float = BODY_FONT_SIZE) -> None:
        for token in tokenize_mixed_content(text):
            if isinstance(token, BlankToken):
                run = paragraph.add_run("\u00a0" * 10)
                _set_run_font(run, size)
                run.underline = True
                continue
            if isinstance(token, str):
                plain = _clean_plain_text(token)
                # 只有当包含小问换行时才保留 \n 软换行，普通单 \n 替换为空格，保持文本紧贴插图流动 (禁用后行断言)
                plain_clean = re.sub(r"\n(?!\s*(?:[\(\（][0-9a-zA-Z一二三四五六七八九十]+\s*[\)\）]|[①-⑩]))", " ", plain)
                chunks = plain_clean.split("\n")
                for index, chunk in enumerate(chunks):
                    if chunk:
                        run = paragraph.add_run(chunk)
                        _set_run_font(run, size)
                    if index < len(chunks) - 1:
                        paragraph.add_run().add_break()
                continue
            formula = token.latex
            xml_bytes = self.omml_map.get(formula)
            if xml_bytes:
                _append_omml(paragraph, xml_bytes)
                self.diagnostics.native_formulas += 1
                continue
            if formula not in self.fallback_cache:
                self.fallback_cache[formula] = _formula_fallback_png(formula)
            image_bytes = self.fallback_cache[formula]
            if image_bytes:
                run = paragraph.add_run()
                try:
                    from PIL import Image

                    image = Image.open(BytesIO(image_bytes))
                    ratio = image.width / max(image.height, 1)
                    height = 0.24 if not token.display else 0.34
                    width = min(6.1, max(0.3, ratio * height))
                    run.add_picture(BytesIO(image_bytes), width=Inches(width), height=Inches(height))
                    self.diagnostics.fallback_formulas += 1
                    continue
                except Exception:
                    pass
            run = paragraph.add_run(f"[公式待核对：{formula}]")
            _set_run_font(run, size, color=RGBColor(192, 0, 0))
            self.diagnostics.failed_formulas += 1

    def add_title_block(
        self,
        title: str,
        subtitle: str,
        total_count: int,
        total_score: int,
        paper_type: str,
        show_secret: bool,
        show_notice: bool,
    ) -> None:
        is_exam_style = paper_type in {"exam", "exam_19"}
        if show_secret and is_exam_style:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run("绝密★启用前")
            _set_run_font(run, 10.5, bold=True, cjk_font=CJK_HEADING_FONT)

        title_p = self.doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(2)
        title_p.paragraph_format.space_after = Pt(3)
        run = title_p.add_run(title or "未命名试卷")
        _set_run_font(run, 18, bold=True, cjk_font=CJK_TITLE_FONT)

        if is_exam_style:
            subject_p = self.doc.add_paragraph()
            subject_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subject_p.paragraph_format.space_before = Pt(2)
            subject_p.paragraph_format.space_after = Pt(4)
            run = subject_p.add_run("数  学")
            _set_run_font(run, 15, bold=True, cjk_font=CJK_TITLE_FONT)

        if subtitle:
            subtitle_p = self.doc.add_paragraph()
            subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_p.paragraph_format.space_before = Pt(2)
            subtitle_p.paragraph_format.space_after = Pt(4)
            run = subtitle_p.add_run(subtitle)
            _set_run_font(run, 12, bold=True, cjk_font=CJK_TITLE_FONT)

        if is_exam_style:
            meta = self.doc.add_paragraph()
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta.paragraph_format.space_before = Pt(4)
            meta.paragraph_format.space_after = Pt(6)
            run = meta.add_run(f"本试卷共 {total_count} 题。全卷满分 {total_score} 分。考试用时 120 分钟。")
            _set_run_font(run, 10.5, cjk_font=CJK_BODY_FONT)

            if show_notice:
                notice_head = self.doc.add_paragraph()
                notice_head.alignment = WD_ALIGN_PARAGRAPH.LEFT
                notice_head.paragraph_format.space_before = Pt(6)
                notice_head.paragraph_format.space_after = Pt(2)
                notice_head.paragraph_format.keep_with_next = True
                run = notice_head.add_run("注意事项：")
                _set_run_font(run, 10.5, bold=True, cjk_font=CJK_HEADING_FONT)

                notices = [
                    "1. 答卷前，考生务必将自己的姓名、考生号、考场号、座位号填写在答题卡上。",
                    "2. 回答选择题时，选出每小题答案后，用铅笔把答题卡上对应题目的答案标号涂黑，如需改动，用橡皮擦干净后，再选涂其他答案标号。回答非选择题时，将答案写在答题卡上。写在本试卷上无效。",
                    "3. 考试结束后，将本试卷和答题卡一并交回。",
                ]
                for idx, line in enumerate(notices):
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Pt(18)
                    p.paragraph_format.first_line_indent = Pt(-18)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(2) if idx < len(notices) - 1 else Pt(6)
                    self.add_mixed(p, line, 10.0)

    def add_section_heading(self, q_type: str, items: list[PreparedQuestion], paper_type: str, chinese_index: int) -> None:
        count = len(items)
        section_score = sum(item.score for item in items)
        scores = {item.score for item in items}
        unit_text = f"每小题 {next(iter(scores))} 分，" if len(scores) == 1 else ""
        label = TYPE_LABELS.get(q_type, "试题")
        if paper_type == "quiz":
            heading = label
        else:
            numerals = "一二三四五六七八九十"
            prefix = numerals[chinese_index - 1] if 0 < chinese_index <= len(numerals) else str(chinese_index)
            heading = f"{prefix}、{label}：本题共 {count} 小题，{unit_text}共 {section_score} 分。"
        p = self.doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        details = {
            "single_choice": "在每小题给出的四个选项中，只有一项符合题目要求。",
            "multi_choice": "在每小题给出的四个选项中，有多项符合题目要求；有选错项不得分。",
            "detailed_answer": "解答应写出文字说明、证明过程或演算步骤。",
        }
        if paper_type != "quiz" and q_type in details:
            heading = f"{heading} {details[q_type]}"

        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(heading)
        _set_run_font(run, SECTION_FONT_SIZE, bold=True, cjk_font=CJK_BODY_FONT)

    def _add_floating_right_image(self, paragraph, path: Path, max_width: float = 1.55) -> bool:
        """
        在 Word 题干段落首行插入原生【四周型右侧文字环绕】插图 (Square Text Wrap)。
        文字可享受 75%+ 的完整宽度空间并于插图左侧自然环绕下流，完全避免无边框表格容器截断。
        """
        try:
            run = paragraph.add_run()
            shape = run.add_picture(str(path), width=Inches(max_width))
            inline = shape._inline
            cx = inline.extent.cx
            cy = inline.extent.cy
            doc_pr = inline.docPr
            doc_id = doc_pr.id
            name = doc_pr.name
            graphic_xml = inline.graphic.xml

            anchor_xml = f'''<wp:anchor {nsdecls('wp', 'r', 'a', 'pic', 'c')}
                distT="36000" distB="36000" distL="144000" distR="36000" simplePos="0" relativeHeight="251658240"
                behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
                <wp:simplePos x="0" y="0"/>
                <wp:positionH relativeFrom="margin">
                    <wp:align>right</wp:align>
                </wp:positionH>
                <wp:positionV relativeFrom="paragraph">
                    <wp:align>top</wp:align>
                </wp:positionV>
                <wp:extent cx="{cx}" cy="{cy}"/>
                <wp:effectExtent l="0" t="0" r="0" b="0"/>
                <wp:wrapSquare wrapText="both"/>
                <wp:docPr id="{doc_id}" name="{name}"/>
                <wp:cNvGraphicFramePr>
                    <a:graphicFrameLocks {nsdecls('a')} noChangeAspect="1"/>
                </wp:cNvGraphicFramePr>
                {graphic_xml}
            </wp:anchor>'''

            anchor_node = parse_xml(anchor_xml)
            inline.getparent().replace(inline, anchor_node)
            return True
        except Exception:
            return False

    def add_question(self, item: PreparedQuestion, number: int) -> None:
        align = item.question.get("figure_align") or "right"
        usable_images = [path for path in item.images if path.exists()]
        if usable_images and align == "right" and len(usable_images) == 1:
            p = self.doc.add_paragraph()
            self._format_question_paragraph(p, number)
            # 在题干段落首行嵌入四周型右侧文字环绕插图，彻底取消表格容器
            self._add_floating_right_image(p, usable_images[0], max_width=1.55)
            # single_paragraph=True 确保全文在同一个 Word 段落中，使文字平滑紧贴图片四周型环绕！
            self._add_content_blocks(p, item.stem, single_paragraph=True)
        else:
            p = self.doc.add_paragraph()
            self._format_question_paragraph(p, number)
            self._add_content_blocks(p, item.stem)
            if usable_images:
                fig_p = self.doc.add_paragraph()
                fig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align == "bottom_right" else WD_ALIGN_PARAGRAPH.CENTER
                for path in usable_images:
                    self._add_image(fig_p, path, max_width=2.35 if len(usable_images) == 1 else 1.85)

        if item.choices:
            self.add_choices(item.choices)

        if item.question.get("question_type") == "detailed_answer" and item.solution_space > 0:
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Cm(max(0.2, item.solution_space - 0.3))

    def _format_question_paragraph(self, paragraph, number: int) -> None:
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.keep_together = True
        run = paragraph.add_run(f"{number}. ")
        _set_run_font(run, BODY_FONT_SIZE, bold=True)

    def _add_content_blocks(self, first_paragraph, content: str, parent=None, single_paragraph: bool = False) -> None:
        parts = re.split(r"\n\s*\n", content.strip()) if content.strip() else [""]
        paragraph = first_paragraph
        for index, part in enumerate(parts):
            if index:
                if single_paragraph:
                    paragraph.add_run().add_break()
                else:
                    paragraph = parent.add_paragraph() if parent is not None else self.doc.add_paragraph()
                    paragraph.paragraph_format.first_line_indent = Cm(0.74)
                    paragraph.paragraph_format.space_after = Pt(2)
            tabular_match = TABULAR_PATTERN.search(part)
            if tabular_match:
                before = _clean_plain_text(part[: tabular_match.start()]).strip()
                after = _clean_plain_text(part[tabular_match.end() :]).strip()
                if before:
                    self.add_mixed(paragraph, before)
                self.add_tabular(tabular_match.group(1), tabular_match.group(2), parent=parent)
                if after:
                    following = paragraph if single_paragraph else (parent.add_paragraph() if parent is not None else self.doc.add_paragraph())
                    self.add_mixed(following, after)
            else:
                self.add_mixed(paragraph, part)

    def _add_image(self, paragraph, path: Path, max_width: float) -> None:
        try:
            run = paragraph.add_run()
            run.add_picture(str(path), width=Inches(max_width))
        except Exception:
            self.diagnostics.missing_images += 1
            self.diagnostics.warnings.append(f"插图无法写入：{path.name}")

    def add_choices(self, choices: list[str]) -> None:
        plain_lengths = [len(re.sub(r"\\[A-Za-z]+|[{}$\\]", "", value)) for value in choices]
        max_length = max(plain_lengths, default=0)
        columns = 4 if len(choices) <= 4 and max_length <= 10 else 2 if max_length <= 24 else 1
        rows = (len(choices) + columns - 1) // columns
        table = self.doc.add_table(rows=rows, cols=columns)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_table_borders(table)
        total_twips = 9000
        width = total_twips // columns
        for row_index, row in enumerate(table.rows):
            _prevent_row_split(row)
            for col_index, cell in enumerate(row.cells):
                _set_cell_width(cell, width)
                _set_cell_margins(cell, 80, 70, 80, 70)
                choice_index = row_index * columns + col_index
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                if choice_index >= len(choices):
                    continue
                label = chr(ord("A") + choice_index)
                run = p.add_run(f"{label}. ")
                _set_run_font(run, BODY_FONT_SIZE, bold=True)
                self.add_mixed(p, choices[choice_index], BODY_FONT_SIZE)
        after = self.doc.add_paragraph()
        after.paragraph_format.space_after = Pt(1)

    def add_tabular(self, column_spec: str, body: str, parent=None) -> None:
        alignments = [char for char in column_spec if char in "lcr"] or ["c"]
        cleaned = re.sub(r"\\(?:toprule|midrule|bottomrule|hline)\b", "", body)
        raw_rows = [row.strip() for row in re.split(r"\\\\(?:\[[^]]*\])?", cleaned) if row.strip()]
        parsed_rows = [_split_table_cells(row) for row in raw_rows]
        if not parsed_rows:
            return
        column_count = max(len(alignments), max(len(row) for row in parsed_rows))
        table = (
            parent.add_table(rows=len(parsed_rows), cols=column_count, width=Inches(4.5))
            if parent is not None
            else self.doc.add_table(rows=len(parsed_rows), cols=column_count)
        )
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_table_borders(table, "7F7F7F", 5)
        width = 9000 // column_count
        for row_index, values in enumerate(parsed_rows):
            row = table.rows[row_index]
            _prevent_row_split(row)
            for col_index, cell in enumerate(row.cells):
                _set_cell_width(cell, width)
                _set_cell_margins(cell, 65, 85, 65, 85)
                if col_index >= len(values):
                    continue
                p = cell.paragraphs[0]
                alignment = alignments[min(col_index, len(alignments) - 1)]
                p.alignment = {"l": WD_ALIGN_PARAGRAPH.LEFT, "c": WD_ALIGN_PARAGRAPH.CENTER, "r": WD_ALIGN_PARAGRAPH.RIGHT}[alignment]
                self.add_mixed(p, values[col_index].strip(), SMALL_FONT_SIZE)
        spacer = parent.add_paragraph() if parent is not None else self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(1)

    def add_answers(self, questions: list[PreparedQuestion], numbers: list[int], title: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.page_break_before = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(14)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(f"{title} 参考答案与解析")
        _set_run_font(run, 16, bold=True, cjk_font=CJK_HEADING_FONT)
        for item, number in zip(questions, numbers):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(f"{number}. ")
            _set_run_font(run, BODY_FONT_SIZE, bold=True)
            if item.answer.strip():
                self._add_content_blocks(p, item.answer)
            else:
                run = p.add_run("暂无答案与解析。")
                _set_run_font(run, SMALL_FONT_SIZE, color=RGBColor(127, 127, 127))

    def save_bytes(self) -> bytes:
        properties = self.doc.core_properties
        properties.title = "MathBank 可编辑 Word 试卷"
        properties.subject = "高中数学试卷"
        properties.author = "MathBank"
        properties.keywords = "数学, 试卷, OMML, 可编辑公式"
        output = BytesIO()
        self.doc.save(output)
        return output.getvalue()


def create_word_bundle_zip(title: str, main_docx_bytes: bytes, ans_docx_bytes: bytes) -> bytes:
    """Create an in-memory ZIP package containing the clean exam DOCX and the answers DOCX."""
    safe_title = re.sub(r'[/\\?%*:|"<>]', "_", title.strip()) or "试卷"
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{safe_title}.docx", main_docx_bytes)
        zf.writestr(f"{safe_title}_含答案与解析.docx", ans_docx_bytes)
    return zip_buffer.getvalue()


def _split_table_cells(row: str) -> list[str]:
    cells: list[str] = []
    current: list[str] = []
    depth = 0
    index = 0
    while index < len(row):
        char = row[index]
        if char == "{" and (index == 0 or row[index - 1] != "\\"):
            depth += 1
        elif char == "}" and (index == 0 or row[index - 1] != "\\"):
            depth = max(0, depth - 1)
        if char == "&" and depth == 0 and (index == 0 or row[index - 1] != "\\"):
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(char)
        index += 1
    cells.append("".join(current).strip())
    return cells


def _resolve_image_paths(question: dict, uploads_dir: str | Path | None) -> list[Path]:
    root = Path(uploads_dir) if uploads_dir else None
    values: list[str] = []
    stored = question.get("image_paths", [])
    if isinstance(stored, list):
        values.extend(str(item) for item in stored)
    values.extend(IMAGE_PATTERN.findall(question.get("content", "") or ""))
    values.extend(IMAGE_PATTERN.findall(question.get("answer_markdown", "") or ""))
    paths: list[Path] = []
    seen: set[str] = set()
    for value in values:
        basename = Path(value).name
        if not basename or basename in seen:
            continue
        seen.add(basename)
        candidate = root / basename if root else Path(value)
        if candidate.exists():
            paths.append(candidate)
    return paths


def _prepare_question(item: dict, uploads_dir: str | Path | None) -> PreparedQuestion:
    question = dict(item.get("question", {}) or {})
    content = question.get("content", "") or ""
    choices: list[str] = []
    match = CHOICES_PATTERN.search(content)
    if match:
        stem = (content[: match.start()] + content[match.end() :]).strip()
        choices = [part.strip() for part in re.split(r"\\item\b", match.group(1)) if part.strip()]
    else:
        stem_lines: list[str] = []
        for line in content.splitlines():
            option = re.match(r"^\s*[-*]?\s*[A-D][\.、]\s*(.*)$", line)
            if option:
                choices.append(option.group(1).strip())
            else:
                stem_lines.append(line)
        stem = "\n".join(stem_lines).strip()
    stem = IMAGE_PATTERN.sub("", stem)
    stem = re.sub(r"\\begin\{tikzpicture\}[\s\S]*?\\end\{tikzpicture\}", "", stem).strip()
    answer = question.get("answer_markdown", "") or ""
    answer_clean = IMAGE_PATTERN.sub("", answer)
    answer_clean = re.sub(r"\\begin\{tikzpicture\}[\s\S]*?\\end\{tikzpicture\}", "", answer_clean).strip()
    try:
        solution_space = max(0.0, float(item.get("solution_space") or question.get("solution_space") or 0))
    except (TypeError, ValueError):
        solution_space = 0.0
    return PreparedQuestion(
        question=question,
        score=int(item.get("score", 5)),
        stem=stem,
        choices=choices,
        answer=answer_clean,
        images=_resolve_image_paths(question, uploads_dir),
        solution_space=solution_space,
    )


def build_word_document(
    title: str,
    subtitle: str,
    paper_type: str,
    questions_data: list,
    *,
    include_answers: bool = False,
    show_secret: bool = True,
    show_notice: bool = True,
    uploads_dir: str | Path | None = None,
) -> tuple[bytes, dict]:
    """Build an A4 editable DOCX and return bytes plus conversion diagnostics."""
    prepared = [_prepare_question(item, uploads_dir) for item in questions_data]
    diagnostics = WordExportDiagnostics(answer_card_omitted=(paper_type == "exam_19"))
    if diagnostics.answer_card_omitted:
        diagnostics.warnings.append("Word 版本仅导出试卷正文，不包含答题卡。")

    formula_sources: list[str] = []
    for item in prepared:
        formula_sources.extend((item.stem, item.answer, *item.choices))
    formulas = _extract_formulas(formula_sources)
    converter = PandocOmmlConverter()
    omml_map = converter.convert_many(formulas)
    if formulas and not converter.pandoc_path:
        diagnostics.warnings.append("未检测到 Pandoc，公式将尝试转为清晰图片。")
    elif len(omml_map) < len(formulas):
        diagnostics.warnings.append("少数公式无法转为 Word 原生公式，已尝试清晰图片兜底。")

    builder = WordExamBuilder(omml_map, diagnostics)
    total_score = sum(item.score for item in prepared)
    builder.add_title_block(
        title,
        subtitle,
        len(prepared),
        total_score,
        paper_type,
        show_secret,
        show_notice,
    )

    grouped: dict[str, list[PreparedQuestion]] = {}
    for item in prepared:
        q_type = item.question.get("question_type", "single_choice")
        grouped.setdefault(q_type, []).append(item)

    number_by_object: dict[int, int] = {}
    next_number = 1
    section_index = 0
    for q_type in TYPE_ORDER:
        items = grouped.get(q_type, [])
        if not items:
            continue
        section_index += 1
        builder.add_section_heading(q_type, items, paper_type, section_index)
        current = EXAM_19_STARTS[q_type] if paper_type == "exam_19" else next_number
        for item in items:
            number_by_object[id(item)] = current
            builder.add_question(item, current)
            current += 1
        if paper_type != "exam_19":
            next_number = current

    if include_answers:
        ordered_numbers = [number_by_object.get(id(item), index + 1) for index, item in enumerate(prepared)]
        builder.add_answers(prepared, ordered_numbers, title)

    output = builder.save_bytes()
    if diagnostics.fallback_formulas:
        diagnostics.warnings.append(f"{diagnostics.fallback_formulas} 处公式以图片保真，不能在 Word 中直接编辑。")
    if diagnostics.failed_formulas:
        diagnostics.warnings.append(f"{diagnostics.failed_formulas} 处公式未能转换，文档中已用红字标出。")
    return output, diagnostics.to_dict()
